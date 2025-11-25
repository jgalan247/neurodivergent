"""
Document Processing utilities for AdaptEd.

Handles extraction of text from PDF, DOCX, and image files.
"""
import io
import logging
from typing import Tuple, Optional
from pathlib import Path

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Processes uploaded documents and extracts text content.
    """

    def process(self, file, file_type: str) -> Tuple[str, int]:
        """
        Process a document and extract text.

        Args:
            file: File object or path
            file_type: Type of file ('pdf', 'docx', 'image')

        Returns:
            Tuple of (extracted_text, page_count)
        """
        if file_type == 'pdf':
            return self._process_pdf(file)
        elif file_type == 'docx':
            return self._process_docx(file)
        elif file_type == 'image':
            return self._process_image(file)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def _process_pdf(self, file) -> Tuple[str, int]:
        """Extract text from a PDF file."""
        try:
            from PyPDF2 import PdfReader

            # Read the file
            if hasattr(file, 'read'):
                pdf_bytes = io.BytesIO(file.read())
                file.seek(0)  # Reset file pointer
            else:
                pdf_bytes = file

            reader = PdfReader(pdf_bytes)
            page_count = len(reader.pages)

            text_parts = []
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"--- Page {i + 1} ---\n{page_text}")

            extracted_text = '\n\n'.join(text_parts)

            logger.info(f"Extracted {len(extracted_text)} characters from {page_count} PDF pages")
            return extracted_text, page_count

        except Exception as e:
            logger.error(f"Error processing PDF: {str(e)}")
            raise

    def _process_docx(self, file) -> Tuple[str, int]:
        """Extract text from a Word document."""
        try:
            from docx import Document

            # Read the file
            if hasattr(file, 'read'):
                doc_bytes = io.BytesIO(file.read())
                file.seek(0)
            else:
                doc_bytes = file

            doc = Document(doc_bytes)

            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)

            extracted_text = '\n\n'.join(text_parts)

            # Estimate page count (roughly 500 words per page)
            word_count = len(extracted_text.split())
            page_count = max(1, word_count // 500)

            logger.info(f"Extracted {len(extracted_text)} characters from DOCX")
            return extracted_text, page_count

        except Exception as e:
            logger.error(f"Error processing DOCX: {str(e)}")
            raise

    def _process_image(self, file) -> Tuple[str, int]:
        """
        Extract text from an image using OCR.

        Note: Requires Tesseract to be installed on the system.
        """
        try:
            from PIL import Image

            # For now, return a placeholder - full OCR would require Tesseract
            # In production, you would use:
            # import pytesseract
            # text = pytesseract.image_to_string(image)

            if hasattr(file, 'read'):
                image = Image.open(file)
                file.seek(0)
            else:
                image = Image.open(file)

            # Placeholder - in production use Tesseract OCR
            logger.warning("OCR not configured - returning placeholder text")
            return "[Image content - OCR processing required]", 1

        except Exception as e:
            logger.error(f"Error processing image: {str(e)}")
            raise


def detect_content_flags(text: str) -> dict:
    """
    Automatically detect content types in the assessment text.

    Args:
        text: The extracted assessment text

    Returns:
        Dictionary of content flags
    """
    import re

    flags = {
        'quotations': False,
        'mathematical_notation': False,
        'code': False,
        'diagrams': False,
        'poetry': False,
        'script': False
    }

    # Check for quotations
    if re.search(r'["""][^"""]+["""]', text) or re.search(r"'[^']{10,}'", text):
        flags['quotations'] = True

    # Check for mathematical notation
    math_patterns = [
        r'\d+\s*[+\-×÷/=<>≤≥]\s*\d+',  # Basic operations
        r'\d+/\d+',  # Fractions
        r'[a-z]\s*[=+\-×÷]\s*\d+',  # Variables
        r'\d+\s*%',  # Percentages
        r'[πθ∑∏∫√]',  # Mathematical symbols
        r'cm²|m²|km²',  # Area units
    ]
    for pattern in math_patterns:
        if re.search(pattern, text, re.IGNORECASE):
            flags['mathematical_notation'] = True
            break

    # Check for code
    code_patterns = [
        r'def\s+\w+\s*\(',  # Python function
        r'function\s+\w+\s*\(',  # JavaScript function
        r'public\s+(?:class|void|static)',  # Java
        r'<[a-z]+[^>]*>',  # HTML tags
        r'\{[^}]*\}',  # Braces (could be code)
        r'print\s*\(',  # Print statements
    ]
    for pattern in code_patterns:
        if re.search(pattern, text):
            flags['code'] = True
            break

    # Check for diagram references
    diagram_patterns = [
        r'(?:figure|fig\.?|diagram|chart|graph|table)\s*\d+',
        r'see\s+(?:the\s+)?(?:figure|diagram|chart)',
        r'shown\s+(?:in|on)\s+(?:the\s+)?(?:figure|diagram)',
    ]
    for pattern in diagram_patterns:
        if re.search(pattern, text, re.IGNORECASE):
            flags['diagrams'] = True
            break

    # Check for poetry (multiple short lines with similar structure)
    lines = text.split('\n')
    short_lines = [l for l in lines if 3 < len(l.split()) < 12]
    if len(short_lines) > 5:
        # Check if many lines have similar lengths (poetry indicator)
        lengths = [len(l) for l in short_lines]
        avg_length = sum(lengths) / len(lengths)
        variance = sum((l - avg_length) ** 2 for l in lengths) / len(lengths)
        if variance < 200:  # Low variance suggests structured verse
            flags['poetry'] = True

    # Check for script/dialogue
    if re.search(r'\b[A-Z][A-Z]+\s*:', text):  # CHARACTER: dialogue format
        flags['script'] = True

    return flags


def suggest_protected_terms(text: str, subject: str) -> list:
    """
    Suggest protected terms based on the content and subject.

    Args:
        text: The assessment text
        subject: The subject of the assessment

    Returns:
        List of suggested protected terms
    """
    import re

    terms = set()

    # Subject-specific term patterns
    subject_patterns = {
        'mathematics': [
            r'\b(?:perimeter|area|volume|radius|diameter|circumference)\b',
            r'\b(?:fraction|numerator|denominator|integer|decimal)\b',
            r'\b(?:equation|formula|variable|coefficient|constant)\b',
            r'\b(?:parallel|perpendicular|adjacent|hypotenuse)\b',
            r'\b(?:mean|median|mode|range|probability)\b',
        ],
        'science': [
            r'\b(?:atom|molecule|electron|proton|neutron)\b',
            r'\b(?:photosynthesis|respiration|osmosis|diffusion)\b',
            r'\b(?:velocity|acceleration|force|momentum|energy)\b',
            r'\b(?:compound|element|mixture|solution)\b',
        ],
        'english': [
            r'\b(?:metaphor|simile|personification|alliteration)\b',
            r'\b(?:protagonist|antagonist|narrator|character)\b',
            r'\b(?:stanza|verse|rhyme|rhythm|imagery)\b',
            r'\b(?:noun|verb|adjective|adverb|pronoun)\b',
        ],
        'history': [
            r'\b(?:century|decade|era|period|dynasty)\b',
            r'\b(?:revolution|reformation|renaissance)\b',
            r'\b(?:monarchy|democracy|republic|empire)\b',
        ],
        'geography': [
            r'\b(?:latitude|longitude|equator|hemisphere)\b',
            r'\b(?:erosion|weathering|deposition|sediment)\b',
            r'\b(?:population|migration|urbanisation)\b',
        ],
    }

    # Get patterns for the subject
    patterns = subject_patterns.get(subject.lower(), [])

    # Find matching terms
    for pattern in patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        terms.update(match.lower() for match in matches)

    # Also find capitalised technical terms
    capitalised = re.findall(r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b', text)
    for term in capitalised:
        if len(term) > 5 and term.lower() not in ['question', 'answer', 'section', 'please']:
            terms.add(term)

    return sorted(list(terms))
